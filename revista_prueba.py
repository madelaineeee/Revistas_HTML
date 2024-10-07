from docx import Document
from bs4 import BeautifulSoup
from tkinter import filedialog as fd
import re
import os

def convert_docx_to_html(input_file, output_file):
    # Load the Word document
    doc = Document(input_file)

    # Create the HTML skeleton with a link to the external CSS
    soup = BeautifulSoup(
        "<html><head><title>New Page</title><link rel='stylesheet' type='text/css' href='main.css'></head><body></body></html>",
        "html.parser")
    body = soup.body

    # Create folder to store images if it doesn't exist
    output_image_folder = os.path.splitext(output_file)[0] + "_images"
    if not os.path.exists(output_image_folder):
        os.makedirs(output_image_folder)

    # Iterate through the paragraphs and other elements in the document
    for para in doc.paragraphs:
        para_text = para.text.strip()

        # Detect numbered headings and special sections
        if re.match(r'^\d+\.\s', para_text):  # Matches "1. ", "2. ", etc.
            p = soup.new_tag('p', **{'class': 'revista_titulo1'})
        elif re.match(r'^\d+\.\d+\s', para_text):  # Matches "1.1 ", "2.1 ", etc.
            p = soup.new_tag('p', **{'class': 'revista_titulo2'})
        elif re.match(r'^\d+\.\d+\.\d+\s', para_text):  # Matches "1.1.1 ", "2.1.1 ", etc.
            p = soup.new_tag('p', **{'class': 'revista_titulo3'})
        elif any(keyword in para_text.upper() for keyword in ["AGRADECIMIENTOS", "CONFLICTO DE INTERESES", "REFERENCIAS"]):
            p = soup.new_tag('p', **{'class': 'revista_titulo1'})
        elif re.match(r'\[\d*?', para_text):  # Matches [1] for references
            p = soup.new_tag('p', **{'class': 'revista_referencias'})
        else:
            p = soup.new_tag('p', **{'class': 'revista_contenido'})  # Default style

        p.string = para_text
        body.append(p)

    # Iterate through tables in the document and convert them to HTML
    for table in doc.tables:
        table_tag = soup.new_tag('table', **{'class': 'revista_tabla'})
        for row in table.rows:
            row_tag = soup.new_tag('tr')
            for cell in row.cells:
                cell_tag = soup.new_tag('td')
                cell_tag.string = cell.text.strip()
                row_tag.append(cell_tag)
            table_tag.append(row_tag)
        body.append(table_tag)

    # Iterate through shapes (images and graphics)
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:  # Detect images
            # Save image to the output folder
            img_data = rel.target_part.blob
            img_filename = f"{output_image_folder}/image_{len(os.listdir(output_image_folder)) + 1}.png"
            with open(img_filename, 'wb') as img_file:
                img_file.write(img_data)

            # Insert the image into the HTML
            img_tag = soup.new_tag('img', src=img_filename)
            body.append(img_tag)

    # Save the HTML to the output file
    with open(output_file, 'w', encoding='utf-8') as file:
        file.write(soup.prettify())

    print(f"Conversion completed successfully. HTML saved as '{output_file}'.")


# Input and output file paths
initialdir = 'C:/Users/50763/Downloads'
input_file = fd.askopenfilename(initialdir=initialdir)
output_file = 'C:/Users/50763/Downloads/articulo.html'

# Convert the file
convert_docx_to_html(input_file, output_file)
print("La conversión se completó con éxito. El archivo HTML se guardó como 'articulo.html'.")
